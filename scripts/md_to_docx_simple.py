from pathlib import Path
import re
from docx import Document


def convert(markdown_path: Path, output_path: Path) -> None:
    text = markdown_path.read_text(encoding="utf-8")
    doc = Document()
    in_code = False

    for raw_line in text.splitlines():
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            doc.add_paragraph(line)
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("!["):
            # Keep image path as text reference in docx.
            doc.add_paragraph(stripped)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        lstripped = line.lstrip()
        if lstripped.startswith("- "):
            doc.add_paragraph(lstripped[2:].strip(), style="List Bullet")
            continue

        if re.match(r"^\d+[\.\)]\s+", stripped):
            item = re.sub(r"^\d+[\.\)]\s+", "", stripped)
            doc.add_paragraph(item, style="List Number")
            continue

        doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    base = Path(r"d:\recipe-hub\recipe-hub")
    md_file = base / "docs" / "SCREENSHOT_INDEX.md"
    docx_file = base / "docs" / "SCREENSHOT_INDEX.docx"
    convert(md_file, docx_file)
    print(docx_file)
